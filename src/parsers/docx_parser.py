import re
import logging
from pathlib import Path
from typing import List, Optional

from docx import Document

from src.models.textbook import Chapter, Textbook
from src.parsers.base import BaseParser
from src.parsers.factory import register_parser

logger = logging.getLogger(__name__)

# 章节标题正则模式
CHAPTER_PATTERNS = [
    r"第[一二三四五六七八九十百千\d]+[章篇]",
    r"Chapter\s+\d+",
    r"CHAPTER\s+\d+",
]


class DocxParser(BaseParser):
    """Word DOCX解析器"""

    def parse(self, file_path: Path) -> Textbook:
        """解析DOCX文件"""
        logger.info(f"Parsing DOCX: {file_path}")
        
        try:
            doc = Document(str(file_path))
            chapters = self._extract_chapters(doc)
            total_chars = sum(ch.char_count for ch in chapters)
            
            return Textbook(
                textbook_id=file_path.stem,
                filename=file_path.name,
                title=self._extract_title_from_filename(file_path),
                total_pages=0,
                total_chars=total_chars,
                chapters=chapters
            )
        except Exception as e:
            logger.error(f"Failed to parse DOCX: {e}")
            raise

    def _extract_chapters(self, doc: Document) -> List[Chapter]:
        """从文档提取章节"""
        chapters: List[Chapter] = []
        
        current_title = ""
        current_content: List[str] = []
        chapter_index = 0
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            if not text:
                continue
            
            # 检测标题样式
            is_heading = self._is_heading(paragraph)
            
            if is_heading:
                # 保存前一章
                if current_content:
                    content = '\n'.join(current_content).strip()
                    if content:
                        chapters.append(Chapter(
                            chapter_id=self._generate_chapter_id(chapter_index),
                            title=current_title or f"章节 {chapter_index + 1}",
                            page_start=0,
                            page_end=0,
                            content=content,
                            char_count=len(content)
                        ))
                        chapter_index += 1
                
                # 开始新章
                current_title = text
                current_content = []
            
            else:
                current_content.append(text)
        
        # 保存最后一章
        if current_content:
            content = '\n'.join(current_content).strip()
            if content:
                chapters.append(Chapter(
                    chapter_id=self._generate_chapter_id(chapter_index),
                    title=current_title or f"章节 {chapter_index + 1}",
                    page_start=0,
                    page_end=0,
                    content=content,
                    char_count=len(content)
                ))
        
        # 如果没有检测到章节，将整个文件作为一章
        if not chapters:
            full_text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
            chapters.append(Chapter(
                chapter_id=self._generate_chapter_id(0),
                title=self._extract_title_from_filename(Path("")),
                page_start=0,
                page_end=0,
                content=full_text,
                char_count=len(full_text)
            ))
        
        return chapters

    def _is_heading(self, paragraph) -> bool:
        """判断段落是否为标题"""
        style_name = paragraph.style.name.lower()
        
        # 检查样式名称
        if 'heading' in style_name or '标题' in style_name:
            return True
        
        # 检查正则模式
        text = paragraph.text.strip()
        for pattern in CHAPTER_PATTERNS:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        
        return False


# 注册解析器
register_parser("docx", DocxParser)
